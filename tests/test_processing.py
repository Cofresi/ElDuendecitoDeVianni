from __future__ import annotations

from copy import copy
from datetime import date
from pathlib import Path

from docx import Document
from openpyxl import Workbook, load_workbook

from el_duendecito_de_vianni.config import AppConfig
from el_duendecito_de_vianni.importer import import_export
from el_duendecito_de_vianni.processed_store import ProcessedStore, file_sha256
from el_duendecito_de_vianni.processor import DocumentProcessor
from el_duendecito_de_vianni.spreadsheet import read_employees
from el_duendecito_de_vianni.spreadsheet import has_employee_rows
from el_duendecito_de_vianni.templates import process_docx, replace_placeholders
from el_duendecito_de_vianni.utils import company_template_subfolder, employee_folder_name, sorted_templates
from el_duendecito_de_vianni.work_schedule import (
    DERIVED_FIELD,
    SchedulePolicy,
    WorkScheduleLookup,
    add_work_schedule_sentence,
    build_work_schedule_sentence,
)
from el_duendecito_de_vianni.mercury import (
    MercuryAutomationError,
    _POST_HR_MARKERS,
    _company_file_label,
    _find_installed_browser,
    _find_playwright_chromium,
    _format_mercury_date,
    _configured_companies,
    _report_generator_url,
    run_mercury_login_test,
)
from el_duendecito_de_vianni.credentials import delete_mercury_password, load_mercury_password, save_mercury_password


def make_config(tmp_path: Path) -> AppConfig:
    return AppConfig(
        downloads_folder=str(tmp_path / "Downloads"),
        template_folder=str(tmp_path / "templates"),
        output_folder=str(tmp_path / "output"),
        imported_folder=str(tmp_path / "imported_files"),
        logs_folder=str(tmp_path / "logs"),
        work_schedule_lookup=str(tmp_path / "politica_horario.xlsx"),
        mercury_url="",
        mercury_username="",
        mercury_company="Supermercado Ines",
        mercury_companies="Supermercado Ines;Brothers",
        mercury_report_name="EntradasDeHoy",
        mercury_headless=False,
    )


def make_employee_sheet(path: Path, rows: list[dict[str, object]]) -> None:
    workbook = Workbook()
    sheet = workbook.active
    headers = list(rows[0].keys())
    sheet.append(headers)
    for row in rows:
        sheet.append([row.get(header) for header in headers])
    workbook.save(path)


def make_docx(path: Path, text: str) -> None:
    document = Document()
    document.add_paragraph(text)
    document.save(path)


def make_schedule_lookup(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["horario1 (de GridViewExport.xls)", "horario2", "dias", "break", "feriados"])
    sheet.append(["02:30 PM A 10:30 PM", "01:30 PM A 8:30 PM", 6, 20, 1])
    sheet.append(["07:00 AM A 03:00 PM", "08:00 AM A 02:00 PM", 5, None, 0])
    workbook.save(path)


def test_one_employee_and_one_docx_template(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    Path(config.output_folder).mkdir(parents=True)
    make_employee_sheet(tmp_path / "employees.xlsx", [{"Numero": 295, "Nombre Empleado": "ALANNA"}])
    make_docx(Path(config.template_folder) / "01_Contrato.docx", "Hola {{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert report.employees_processed == 1
    assert report.document_count == 1
    generated = Path(report.generated_documents[0])
    assert generated.exists()
    assert "Hola ALANNA" in "\n".join(p.text for p in Document(generated).paragraphs)


def test_multiple_employees_and_multiple_templates(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [
            {"Numero": 1, "Nombre Empleado": "ANA"},
            {"Numero": 2, "Nombre Empleado": "LUIS"},
        ],
    )
    make_docx(Path(config.template_folder) / "01_Contrato.docx", "{{Nombre Empleado}}")
    make_docx(Path(config.template_folder) / "02_Autorizacion.docx", "{{Numero}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert report.employees_processed == 2
    assert report.document_count == 4


def test_static_template_is_copied_without_placeholder_processing(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(tmp_path / "employees.xlsx", [{"Numero": 1, "Nombre Empleado": "ANA"}])
    make_docx(
        Path(config.template_folder) / "03__STATIC__Reglamento.docx",
        "Texto fijo {{Nombre Empleado}}",
    )

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert report.document_count == 1
    generated = Path(report.generated_documents[0])
    assert "Texto fijo {{Nombre Empleado}}" in "\n".join(p.text for p in Document(generated).paragraphs)
    assert not report.skipped_files


def test_placeholders_with_accents_and_spaces(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Posición": "Analista"}],
    )
    make_docx(Path(config.template_folder) / "01_Contrato.docx", "Cargo: {{Posición}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    text = "\n".join(p.text for p in Document(report.generated_documents[0]).paragraphs)
    assert "Cargo: Analista" in text


def test_blank_excel_cells_become_blank_text(tmp_path: Path) -> None:
    path = tmp_path / "employees.xlsx"
    make_employee_sheet(path, [{"Numero": 1, "Nombre Empleado": None, "Telefono1": None}])

    rows = read_employees(path)

    assert rows[0]["Nombre Empleado"] == ""
    assert rows[0]["Telefono1"] == ""


def test_empty_employee_sheet_has_no_employee_rows(tmp_path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["Numero", "Nombre Empleado"])
    path = tmp_path / "empty.xlsx"
    workbook.save(path)

    assert not has_employee_rows(path)


def test_management_tile_label_is_not_hr_page_marker() -> None:
    assert "Recursos Humanos" not in _POST_HR_MARKERS


def test_mercury_requires_configured_url(tmp_path: Path) -> None:
    config = make_config(tmp_path)

    try:
        run_mercury_login_test(config, "secret")
    except MercuryAutomationError as exc:
        assert "direccion de Mercury" in str(exc)
    else:
        raise AssertionError("Expected MercuryAutomationError")


def test_mercury_requires_saved_password(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    config.mercury_url = "https://example.test"
    config.mercury_username = "usuario"

    try:
        run_mercury_login_test(config, "")
    except MercuryAutomationError as exc:
        assert "contrasena de Mercury" in str(exc)
    else:
        raise AssertionError("Expected MercuryAutomationError")


def test_find_playwright_chromium_from_local_app_data(tmp_path: Path, monkeypatch) -> None:
    chrome = tmp_path / "ms-playwright" / "chromium-1228" / "chrome-win64" / "chrome.exe"
    chrome.parent.mkdir(parents=True)
    chrome.write_text("fake browser", encoding="utf-8")
    monkeypatch.setenv("LOCALAPPDATA", str(tmp_path))
    monkeypatch.delenv("EL_DUENDECITO_CHROMIUM_EXE", raising=False)

    assert _find_playwright_chromium() == chrome


def test_find_installed_edge_when_playwright_browser_is_missing(tmp_path: Path, monkeypatch) -> None:
    edge = tmp_path / "Microsoft" / "Edge" / "Application" / "msedge.exe"
    edge.parent.mkdir(parents=True)
    edge.write_text("fake edge", encoding="utf-8")
    monkeypatch.setenv("PROGRAMFILES", str(tmp_path))
    monkeypatch.delenv("PROGRAMFILES(X86)", raising=False)
    monkeypatch.delenv("LOCALAPPDATA", raising=False)

    assert _find_installed_browser() == edge


def test_report_generator_url_uses_mercury_host() -> None:
    assert (
        _report_generator_url("http://192.168.1.3/Mercury.Menu/Management.aspx?id=abc")
        == "http://192.168.1.3/Mercury.RRHH/GeneradorReportes.aspx"
    )


def test_mercury_date_format_for_filters() -> None:
    assert _format_mercury_date(date(2026, 7, 6)) == "06.07.2026"


def test_mercury_company_list_and_file_labels(tmp_path: Path) -> None:
    config = make_config(tmp_path)

    assert _configured_companies(config) == ["Supermercado Ines", "Brothers"]
    assert _company_file_label("Supermercado Ines") == "Ines"
    assert _company_file_label("Brothers") == "Brothers"


def test_mercury_password_round_trip(tmp_path: Path, monkeypatch) -> None:
    monkeypatch.setenv("EL_DUENDECITO_CREDENTIAL_DIR", str(tmp_path))

    save_mercury_password("clave-secreta")

    assert load_mercury_password() == "clave-secreta"
    delete_mercury_password()
    assert load_mercury_password() == ""


def test_missing_placeholder_is_left_and_reported(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(tmp_path / "employees.xlsx", [{"Numero": 1, "Nombre Empleado": "ANA"}])
    make_docx(Path(config.template_folder) / "01_Contrato.docx", "{{Campo Inexistente}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert "Campo Inexistente" in report.missing_placeholders
    text = "\n".join(p.text for p in Document(report.generated_documents[0]).paragraphs)
    assert "{{Campo Inexistente}}" in text


def test_placeholder_split_across_word_runs(tmp_path: Path) -> None:
    path = tmp_path / "split.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Hola {{Nombre")
    paragraph.add_run(" Empleado}}")
    document.save(path)

    missing: set[str] = set()
    process_docx(path, {"Nombre Empleado": "ANA"}, missing)

    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert text == "Hola ANA"
    assert not missing


def test_placeholder_format_filters() -> None:
    missing: set[str] = set()
    values = {
        "Salario Base": "18800",
        "Numero": "295.0",
        "Fecha Ingreso": "2026-07-04",
        "Nombre Empleado": "ANA",
        "Sexo": "Femenino",
    }

    rendered = replace_placeholders(
        "{{Salario Base|money}} {{Numero|int}} {{Fecha Ingreso|date}} {{Sexo|tratamiento}} {{Nombre Empleado}}",
        values,
        missing,
    )

    assert rendered == "18,800.00 295 04/07/2026 Sra. ANA"
    assert not missing


def test_tratamiento_filter_gender_values() -> None:
    assert replace_placeholders("{{Sexo|tratamiento}}", {"Sexo": "F"}, set()) == "Sra."
    assert replace_placeholders("{{Sexo|tratamiento}}", {"Sexo": "Femenino"}, set()) == "Sra."
    assert replace_placeholders("{{Sexo|tratamiento}}", {"Sexo": "M"}, set()) == "Sr."
    assert replace_placeholders("{{Sexo|tratamiento}}", {"Sexo": "Masculino"}, set()) == "Sr."
    assert replace_placeholders("{{Sexo|tratamiento}}", {"Sexo": "Sin definir"}, set()) == ""


def test_genero_ending_filters() -> None:
    assert replace_placeholders("Estimad{{Sexo|genero}}", {"Sexo": "Femenino"}, set()) == "Estimada"
    assert replace_placeholders("Estimad{{Sexo|genero}}", {"Sexo": "Masculino"}, set()) == "Estimado"
    assert replace_placeholders("Estimad{{Sexo|genero_plural}}", {"Sexo": "F"}, set()) == "Estimadas"
    assert replace_placeholders("Estimad{{Sexo|genero_plural}}", {"Sexo": "M"}, set()) == "Estimados"
    assert replace_placeholders("Estimad{{Sexo|genero}}", {"Sexo": ""}, set()) == "Estimad"


def test_genero_sustantivo_ending_filters() -> None:
    assert replace_placeholders("colaborador{{Sexo|genero_sustantivo}}", {"Sexo": "F"}, set()) == "colaboradora"
    assert replace_placeholders("colaborador{{Sexo|genero_sustantivo}}", {"Sexo": "M"}, set()) == "colaborador"
    assert replace_placeholders("colaborador{{Sexo|genero_sustantivo_plural}}", {"Sexo": "Femenino"}, set()) == "colaboradoras"
    assert replace_placeholders("colaborador{{Sexo|genero_sustantivo_plural}}", {"Sexo": "Masculino"}, set()) == "colaboradores"
    assert replace_placeholders("colaborador{{Sexo|genero_sustantivo}}", {"Sexo": ""}, set()) == "colaborador"


def test_work_schedule_sentence_for_six_day_policy() -> None:
    sentence = build_work_schedule_sentence(
        SchedulePolicy(
            horario1="02:30 PM A 10:30 PM",
            horario2="01:30 PM A 8:30 PM",
            days=6,
            break_minutes=20,
            holidays=True,
        )
    )

    assert (
        sentence
        == "Lunes a Sábados de 2:30 pm a 10:30 pm y Domingo de 1:30 pm a 8:30 pm, "
        "con 20 minutos de break y un día libre a la semana según programación, "
        "este horario incluye los días feriados."
    )


def test_work_schedule_sentence_omits_holidays_when_disabled() -> None:
    sentence = build_work_schedule_sentence(
        SchedulePolicy(
            horario1="07:00 AM A 03:00 PM",
            horario2="08:00 AM A 02:00 PM",
            days=5,
            break_minutes=None,
            holidays=False,
        )
    )

    assert sentence == "Lunes a Viernes de 7:00 am a 3:00 pm y Sábado de 8:00 am a 2:00 pm."


def test_work_schedule_lookup_enriches_employee(tmp_path: Path) -> None:
    lookup_path = tmp_path / "politica_horario.xlsx"
    make_schedule_lookup(lookup_path)
    employee = {"Política Horario": "02:30 PM A 10:30 PM"}

    add_work_schedule_sentence(employee, WorkScheduleLookup.from_file(lookup_path))

    assert employee[DERIVED_FIELD].startswith("Lunes a Sábados de 2:30 pm")


def test_unknown_work_schedule_policy_becomes_blank(tmp_path: Path) -> None:
    lookup_path = tmp_path / "politica_horario.xlsx"
    make_schedule_lookup(lookup_path)
    employee = {"Política Horario": "No existe"}

    add_work_schedule_sentence(employee, WorkScheduleLookup.from_file(lookup_path))

    assert employee[DERIVED_FIELD] == ""


def test_docx_template_money_filter(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Salario Base": 18800}],
    )
    make_docx(Path(config.template_folder) / "01_Contrato.docx", "Salario: {{Salario Base|money}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    text = "\n".join(p.text for p in Document(report.generated_documents[0]).paragraphs)
    assert "Salario: 18,800.00" in text


def test_docx_template_tratamiento_filter(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Sexo": "Femenino"}],
    )
    make_docx(Path(config.template_folder) / "01_Carta.docx", "{{Sexo|tratamiento}} {{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    text = "\n".join(p.text for p in Document(report.generated_documents[0]).paragraphs)
    assert "Sra. ANA" in text


def test_docx_template_work_schedule_field(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_schedule_lookup(Path(config.work_schedule_lookup))
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Política Horario": "02:30 PM A 10:30 PM"}],
    )
    make_docx(Path(config.template_folder) / "01_Carta.docx", "{{Horario Laboral}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    text = "\n".join(p.text for p in Document(report.generated_documents[0]).paragraphs)
    assert "Lunes a Sábados de 2:30 pm a 10:30 pm" in text


def test_xlsx_template_preserves_formula_and_formatting(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(tmp_path / "employees.xlsx", [{"Numero": 1, "Nombre Empleado": "ANA"}])
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "{{Nombre Empleado}}"
    sheet["B1"] = "=1+1"
    bold_font = copy(sheet["A1"].font)
    bold_font.bold = True
    sheet["A1"].font = bold_font
    template = Path(config.template_folder) / "01_Formulario.xlsx"
    workbook.save(template)

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")
    generated = load_workbook(report.generated_documents[0], data_only=False)

    assert generated.active["A1"].value == "ANA"
    assert generated.active["A1"].font.bold
    assert generated.active["B1"].value == "=1+1"


def test_xlsx_template_format_filters(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    Path(config.template_folder).mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 295, "Nombre Empleado": "ANA", "Salario Base": 18800}],
    )
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = "{{Salario Base|money}}"
    sheet["A2"] = "{{Numero|int}}"
    template = Path(config.template_folder) / "01_Formulario.xlsx"
    workbook.save(template)

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")
    generated = load_workbook(report.generated_documents[0])

    assert generated.active["A1"].value == "18,800.00"
    assert generated.active["A2"].value == "295"


def test_duplicate_input_export_detection(tmp_path: Path) -> None:
    downloads = tmp_path / "Downloads"
    imported = tmp_path / "imported_files"
    downloads.mkdir()
    source = downloads / "GridViewExports.xlsx"
    make_employee_sheet(source, [{"Numero": 1, "Nombre Empleado": "ANA"}])
    result = import_export(source, imported)
    store = ProcessedStore(imported / "processed_files.json")
    store.add(source, result.imported_path, file_sha256(result.imported_path))

    assert store.is_processed_hash(file_sha256(source))


def test_invalid_characters_in_employee_names() -> None:
    row = {"Numero": "12", "Nombre Empleado": 'ANA:LUIS/TEST*"'}

    assert employee_folder_name(row) == "12 - ANA-LUIS-TEST--"


def test_print_order_sorting() -> None:
    folder = Path("unused")
    names = ["10_Final.docx", "02_Formulario.xlsx", "01_Contrato.docx"]
    assert sorted(names, key=lambda value: value.lower()) == ["01_Contrato.docx", "02_Formulario.xlsx", "10_Final.docx"]


def test_template_sorting_ignores_office_lock_files(tmp_path: Path) -> None:
    make_docx(tmp_path / "01_Contrato.docx", "Contrato")
    (tmp_path / "~$01_Contrato.docx").write_text("office lock", encoding="utf-8")
    (tmp_path / "~$02_Formulario.xlsx").write_text("office lock", encoding="utf-8")

    assert [path.name for path in sorted_templates(tmp_path)] == ["01_Contrato.docx"]


def test_company_template_routing_uses_ines_folder(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    ines_folder = Path(config.template_folder) / "Ines"
    brothers_folder = Path(config.template_folder) / "Brothers"
    ines_folder.mkdir(parents=True)
    brothers_folder.mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Compania": "Supermercado Ines"}],
    )
    make_docx(ines_folder / "01_Ines.docx", "Ines {{Nombre Empleado}}")
    make_docx(brothers_folder / "01_Brothers.docx", "Brothers {{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert report.document_count == 1
    generated = Path(report.generated_documents[0])
    assert generated.name == "01_Ines.docx"
    assert "Ines ANA" in "\n".join(p.text for p in Document(generated).paragraphs)
    assert Path(report.output_folder).parent.name.startswith("nuevasEntradas_")
    assert Path(report.output_folder).name == "Ines"


def test_selected_run_date_controls_output_folder(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    ines_folder = Path(config.template_folder) / "Ines"
    ines_folder.mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Compania": "Supermercado Ines"}],
    )
    make_docx(ines_folder / "01_Ines.docx", "Ines {{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx", run_date=date(2026, 7, 6))

    assert Path(report.output_folder).parent.name == "nuevasEntradas_06.07.2026"


def test_company_template_routing_uses_brothers_for_other_companies(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    ines_folder = Path(config.template_folder) / "Ines"
    brothers_folder = Path(config.template_folder) / "Brothers"
    ines_folder.mkdir(parents=True)
    brothers_folder.mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 2, "Nombre Empleado": "LUIS", "Compania": "Otra Compania"}],
    )
    make_docx(ines_folder / "01_Ines.docx", "Ines {{Nombre Empleado}}")
    make_docx(brothers_folder / "01_Brothers.docx", "Brothers {{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert report.document_count == 1
    generated = Path(report.generated_documents[0])
    assert generated.name == "01_Brothers.docx"
    assert "Brothers LUIS" in "\n".join(p.text for p in Document(generated).paragraphs)
    assert Path(report.output_folder).parent.name.startswith("nuevasEntradas_")
    assert Path(report.output_folder).name == "Brothers"


def test_company_template_subfolder_mapping() -> None:
    assert company_template_subfolder({"Compania": "Supermercado Ines"}) == "Ines"
    assert company_template_subfolder({"Compania": "Supermercado Inés"}) == "Brothers"
    assert company_template_subfolder({"Compania": ""}) == "Brothers"


def test_processing_message_includes_company_name(tmp_path: Path) -> None:
    config = make_config(tmp_path)
    ines_folder = Path(config.template_folder) / "Ines"
    ines_folder.mkdir(parents=True)
    make_employee_sheet(
        tmp_path / "employees.xlsx",
        [{"Numero": 1, "Nombre Empleado": "ANA", "Compania": "Supermercado Ines"}],
    )
    make_docx(ines_folder / "01_Ines.docx", "{{Nombre Empleado}}")

    report = DocumentProcessor(config).process_imported_file(tmp_path / "employees.xlsx")

    assert "nuevos empleados de Supermercado Ines" in report.message
